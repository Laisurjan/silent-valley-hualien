from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

ROOT = Path(__file__).resolve().parent.parent
OUT = ROOT / "teaching_materials" / "student_report_template.docx"

EVENTS = [
    ("cepo", "大港口事件 Cepo'", "1877-1878", "阿美族"),
    ("dafen", "大分事件", "1914-1933", "布農族"),
    ("cikasuan", "七腳川事件", "1908-1914", "阿美族"),
    ("truku", "太魯閣戰役", "1914", "太魯閣族"),
]

NOTEBOOK_QUESTIONS = [
    "大港口事件：請根據 PDF，整理清軍、地方官員與 Cepo' 部落對事件原因的三種說法。",
    "大分事件：請整理 1914-1933 年布農族抵抗、遷移與 Dahu Ali 相關線索。",
    "七腳川事件：請說明七腳川社與日本官方說法的差異，並列出可引用頁碼。",
    "太魯閣戰役：請比較日方理蕃政策與太魯閣族守護土地的觀點。",
]

STANCE_TABLE_INTRO = "漢/日方與族方對同一事件常用不同詞彙。AI 答案多偏漢/日視角,你的工作是找出至少一個並改寫成更中性或族視角的版本。"
STANCE_TABLE_HEADERS = ["詞彙", "漢/日視角", "族視角", "適用事件", "過度簡化風險"]
STANCE_TABLE_ROWS = [
    ["招撫", "招撫原住民、撫番", "國家力量介入、要求服從", "大港口", "不宜一律譯成「強迫歸順」"],
    ["平亂", "官兵殺平、攻剿", "抵抗入侵後遭鎮壓", "大港口", "可改「鎮壓抵抗」,保留清廷視為亂事的立場"],
    ["滅社", "喪身滅社、討伐滅社", "部落被拆散、遷徙、失去故土", "大港口、七腳川", "「清軍屠殺」只適大港口;七腳川是日治「討伐滅社」"],
    ["理蕃", "治理蕃地的政策", "殖民管控、分類與改造", "大分、七腳川、太魯閣", "標出日治原詞"],
    ["歸順", "臣服、投降、歸順日本", "被迫投降或戰術性和解", "大分、七腳川、太魯閣", "「被迫投降」太單一;大分有 Dahu Ali「和解」與「最後未歸順蕃」的差別"],
    ["收押槍枝", "沒收以便治理", "剝奪生計、防衛、狩獵能力", "大分、太魯閣", "大分事件導火線之一"],
    ["駐在所", "山地警察治理據點", "壓迫、監控、衝突目標", "大分、太魯閣", "大分多次涉襲擊駐在所"],
    ["討伐／征討", "軍警合法征伐", "國家級武力入侵", "七腳川、太魯閣", "比「平亂」更日治準確"],
    ["隘勇線", "警備防線、理蕃設施", "壓縮生活領域、迫使衝突", "七腳川、太魯閣", "七腳川叢書目次列「協助隘勇線的建立」"],
    ["太魯閣蕃討伐", "日方戰役稱呼", "太魯閣族抗日／保衛土地", "太魯閣", "詞彙本身就是立場詞"],
    ["正名", "從事件改稱戰役", "歷史平反、恢復族群敘事", "大港口", "2022 正名 Cepo' 戰役"],
]


def set_font(run, size=11, bold=False, color=None):
    run.font.name = "Microsoft JhengHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft JhengHei")
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_width(cell, width_cm):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_w = tc_pr.first_child_found_in("w:tcW")
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(int(width_cm * 567)))
    tc_w.set(qn("w:type"), "dxa")


def p(cell, text="", size=10.5, bold=False):
    para = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
    para.text = ""
    run = para.add_run(text)
    set_font(run, size=size, bold=bold)
    return para


def add_box(doc, title, rows):
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    table.autofit = True
    table.cell(0, 0).merge(table.cell(0, 1))
    shade(table.cell(0, 0), "EDE6D6")
    p(table.cell(0, 0), title, 12, True)
    for label, hint in rows:
        cells = table.add_row().cells
        set_width(cells[0], 4.2)
        shade(cells[0], "F7F3EA")
        p(cells[0], label, 10, True)
        p(cells[1], hint, 10)
        cells[0].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
        cells[1].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
    doc.add_paragraph()
    return table


def add_text_line(doc, text, size=10):
    para = doc.add_paragraph()
    run = para.add_run(text)
    set_font(run, size=size)
    return para


def add_stance_table(doc):
    add_text_line(doc, "立場詞對照表 ─ 找出 AI 用了哪一邊的詞", 12)
    add_text_line(doc, STANCE_TABLE_INTRO, 9.5)
    table = doc.add_table(rows=1, cols=5)
    table.style = "Table Grid"
    table.autofit = True
    for i, text in enumerate(STANCE_TABLE_HEADERS):
        cell = table.cell(0, i)
        shade(cell, "EDE6D6")
        p(cell, text, 8.5, True)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
    for r_idx, row_data in enumerate(STANCE_TABLE_ROWS, start=1):
        cells = table.add_row().cells
        fill = "FFFBF2" if r_idx % 2 else "F0E9DA"
        for c_idx, text in enumerate(row_data):
            shade(cells[c_idx], fill)
            p(cells[c_idx], text, 8)
            cells[c_idx].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
    doc.add_paragraph()
    return table


def build():
    doc = Document()
    sec = doc.sections[0]
    sec.orientation = WD_ORIENT.PORTRAIT
    sec.top_margin = Cm(1.4)
    sec.bottom_margin = Cm(1.4)
    sec.left_margin = Cm(1.4)
    sec.right_margin = Cm(1.4)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("縱谷無言．事件理解卡")
    set_font(run, size=18, bold=True)
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(sub.add_run("單張卡片版｜請用黑筆或深色字填寫，方便影印與教師批閱"), size=10)

    info = doc.add_table(rows=1, cols=6)
    info.style = "Table Grid"
    labels = ["班級", "", "座號", "", "姓名", ""]
    for i, text in enumerate(labels):
        p(info.cell(0, i), text, 10, i % 2 == 0)
        if i % 2 == 0:
            shade(info.cell(0, i), "F0E9DA")
    doc.add_paragraph()

    add_box(doc, "1. 我研究的事件（勾選一項）", [
        ("事件選擇", "□ 大港口事件 Cepo'　□ 大分事件　□ 七腳川事件　□ 太魯閣戰役"),
    ])

    add_box(doc, "2. 事件基礎資料", [
        ("時間", ""),
        ("地點", ""),
        ("族群", ""),
        ("人物", ""),
    ])

    add_text_line(doc, "NotebookLM 開場問題參考", 11)
    for q in NOTEBOOK_QUESTIONS:
        add_text_line(doc, "□ " + q, 9.5)
    add_box(doc, "3. AI 協助我理解", [
        ("我問 NotebookLM 的問題", ""),
        ("NotebookLM 給我的 3 個重點", "1.\n2.\n3."),
        ("我抓到的 3 個關鍵詞", "1.\n2.\n3."),
    ])

    add_box(doc, "4. 我改掉 AI 一個地方", [
        ("AI 原句", ""),
        ("我的改寫", ""),
        ("我修改的理由", ""),
        ("視覺提醒", "把 AI 立場詞劃線或刪除，再用括號標出你換上的較精準詞。例：平定／叛亂 → 鎮壓／守護土地。"),
    ])

    add_stance_table(doc)

    add_box(doc, "5. 立場差異", [
        ("官方說法", ""),
        ("部落／族人說法", ""),
        ("最大的差異", ""),
    ])

    add_box(doc, "6. 地景與記憶", [
        ("地景類型", "□ 河流／溪流　□ 部落／舊社　□ 道路／古道　□ 紀念碑／遺址"),
        ("現在看起來像什麼", ""),
        ("知道歷史後，我會怎麼看", ""),
    ])

    add_box(doc, "7. 三句展示句", [
        ("事實句", "我知道這件事發生在……"),
        ("差異句", "不同立場的差異是……"),
        ("省思句", "我重新理解花蓮地景是……"),
    ])

    add_box(doc, "8. 資料來源", [
        ("來源 1", "標題／頁碼："),
        ("來源 2", "標題／頁碼："),
        ("來源 3", "網址或補充頁碼："),
    ])

    add_box(doc, "9. 同儕互評：一勾一改", [
        ("三個勾選", "□ 事實句有時間、地點或人物　□ 差異句有兩種立場　□ 省思句不是只說很可憐"),
        ("我建議同學改成這一句", ""),
    ])

    add_box(doc, "下半部：AI process log", [
        ("我怎麼判斷 AI 哪裡需要改", ""),
        ("我最後保留／刪除／查證了什麼", ""),
    ])

    add_box(doc, "右側評分軌道（0-2 分）", [
        ("基本事實", "0　1　2"),
        ("AI 使用", "0　1　2"),
        ("立場比較", "0　1　2"),
        ("展示句", "0　1　2"),
    ])

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(f"wrote {OUT.relative_to(ROOT)}")


if __name__ == "__main__":
    build()
