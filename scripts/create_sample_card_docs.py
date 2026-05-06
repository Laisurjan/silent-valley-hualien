from pathlib import Path

from docx import Document

ROOT = Path(__file__).resolve().parent.parent
OUT = ROOT / "data" / "submissions"


def make(path, class_name, seat, name, event_label, fact, difference, reflection):
    doc = Document()
    doc.add_heading("縱谷無言．事件理解卡", 0)
    info = doc.add_table(rows=1, cols=6)
    info.style = "Table Grid"
    for i, value in enumerate(["班級", class_name, "座號", seat, "姓名", name]):
        info.cell(0, i).text = value

    rows = [
        ("事件選擇", event_label),
        ("時間", "1908-1914"),
        ("地點", "花蓮七腳川溪流域"),
        ("族群", "阿美族"),
        ("人物", "七腳川社族人、日本官方"),
        ("我問 NotebookLM 的問題", "請比較官方記載與部落觀點，並列出可引用頁碼。"),
        ("NotebookLM 給我的 3 個重點", "1. 官方重視治安與控制\n2. 部落重視土地與生存\n3. 事件造成遷徙"),
        ("我抓到的 3 個關鍵詞", "七腳川社\n鎮壓\n遷徙"),
        ("AI 原句", "日本政府平定了七腳川社的叛亂。"),
        ("我的改寫", "日本政府以武力鎮壓七腳川社，族人則從保護部落與土地理解這場衝突。"),
        ("我修改的理由", "平定和叛亂只呈現官方立場。"),
        ("官方說法", "維持秩序、處理抗命。"),
        ("部落／族人說法", "保護部落土地與生活。"),
        ("最大的差異", difference),
        ("地景類型", "河流／溪流\n部落／舊社"),
        ("現在看起來像什麼", "一般溪流與聚落。"),
        ("知道歷史後，我會怎麼看", "地名背後有被迫遷徙與記憶。"),
        ("事實句", fact),
        ("差異句", difference),
        ("省思句", reflection),
        ("來源 1", "原民會叢書 PDF，P.23"),
        ("來源 2", "國教院補充教材 PDF，P.5"),
        ("來源 3", "https://storymaps.arcgis.com/collections/7e917133a47a4548bc33f5138a397474"),
        ("三個勾選", "事實句有時間、地點或人物\n差異句有兩種立場\n省思句不是只說很可憐"),
        ("我建議同學改成這一句", "把地點補進事實句。"),
        ("我怎麼判斷 AI 哪裡需要改", "我檢查 AI 是否只用了官方立場詞。"),
        ("我最後保留／刪除／查證了什麼", "保留時間，刪除平定，查證頁碼。"),
        ("基本事實", "2"),
        ("AI 使用", "2"),
        ("立場比較", "2"),
        ("展示句", "2"),
    ]
    for label, value in rows:
        table = doc.add_table(rows=1, cols=2)
        table.style = "Table Grid"
        table.cell(0, 0).text = label
        table.cell(0, 1).text = value

    OUT.mkdir(parents=True, exist_ok=True)
    doc.save(OUT / path)


make(
    "sample_card_01.docx",
    "商二甲",
    "01",
    "測試甲",
    "七腳川事件",
    "我知道七腳川事件發生在 1908 年後的花蓮七腳川溪流域。",
    "官方說法強調治理秩序，部落觀點強調土地與生存被壓迫。",
    "我重新理解花蓮的溪流不只是地景，也可能保存族人遷徙與受傷的記憶。",
)
make(
    "sample_card_02.docx",
    "商二甲",
    "02",
    "測試乙",
    "太魯閣戰役",
    "我知道太魯閣戰役發生在 1914 年，與太魯閣族和日本理蕃政策有關。",
    "官方說法看見統治推進，族人說法看見守護家園。",
    "我重新理解山路和古道不只是觀光路線，也連著族人保護土地的故事。",
)
print("wrote 2 sample card docx files")
