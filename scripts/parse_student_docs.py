from __future__ import annotations

import json
import re
import sys
from datetime import datetime
from pathlib import Path

from docx import Document

ROOT = Path(__file__).resolve().parent.parent
SUBMIT_DIR = ROOT / "data" / "submissions"
OUTPUT_JSON = ROOT / "data" / "class_data.json"

EVENT_KEYWORDS = {
    "cepo": ["大港口", "Cepo"],
    "dafen": ["大分"],
    "cikasuan": ["七腳川", "Cikasuan"],
    "truku": ["太魯閣"],
}


def clean(text: str) -> str:
    return re.sub(r"\n{3,}", "\n\n", str(text or "").strip())


def split_lines(text: str) -> list[str]:
    lines = []
    for raw in re.split(r"[\n；;]+", text or ""):
        item = re.sub(r"^\s*(?:\d+[\.\、)]|[-•□☑✓])\s*", "", raw).strip()
        if item:
            lines.append(item)
    return lines


def detect_event(text: str) -> str:
    for code, keys in EVENT_KEYWORDS.items():
        if any(k in text for k in keys):
            return code
    return ""


def iter_table_pairs(doc: Document) -> dict[str, list[str]]:
    pairs: dict[str, list[str]] = {}
    for table in doc.tables:
        for row in table.rows:
            cells = [clean(c.text) for c in row.cells]
            if not cells:
                continue
            if len(cells) >= 6 and "班級" in cells[0] and "座號" in cells[2] and "姓名" in cells[4]:
                pairs.setdefault("班級", []).append(cells[1])
                pairs.setdefault("座號", []).append(cells[3])
                pairs.setdefault("姓名", []).append(cells[5])
                continue
            if len(cells) >= 2 and cells[0]:
                label = cells[0].replace("　", "").strip()
                value = "\n".join(c for c in cells[1:] if c and c != cells[0]).strip()
                if label:
                    pairs.setdefault(label, []).append(value)
    return pairs


def first(pairs: dict[str, list[str]], *labels: str) -> str:
    for label in labels:
        for key, vals in pairs.items():
            if label in key and vals:
                return clean(vals[0])
    return ""


def score_value(raw: str):
    m = re.search(r"[0-2]", raw or "")
    return int(m.group(0)) if m else None


def parse_sources(*raw_sources: str) -> list[dict]:
    out = []
    for raw in raw_sources:
        raw = clean(raw)
        if not raw:
            continue
        pdf_page = ""
        page_match = re.search(r"(?:p\.?|P\.?|頁|第)\s*([0-9０-９\-–—]+)", raw)
        if page_match:
            pdf_page = page_match.group(0)
        url_match = re.search(r"https?://\S+", raw)
        url = url_match.group(0) if url_match else ""
        title = raw.replace(url, "").strip(" \n：:")
        out.append({
            "type": "NotebookLM" if "NotebookLM" in raw else ("PDF" if "PDF" in raw or "pdf" in raw else ""),
            "title": title,
            "url": url,
            "pdf_page": pdf_page,
        })
    return out


def build_id(class_name: str, seat: str, name: str, fallback: str) -> str:
    parts = [re.sub(r"\s+", "", p) for p in (class_name, seat.zfill(2) if seat else "", name) if p]
    if not parts:
        return re.sub(r"[^A-Za-z0-9_\-\u4e00-\u9fff]+", "_", fallback)
    return "_".join(parts)


def legacy_supplement_if_present(text: str, pairs: dict[str, list[str]]) -> dict:
    old_markers = ["事件經過", "關鍵圖片", "多視角", "個人省思", "事件影響", "第三頁"]
    if not any(m in text for m in old_markers):
        return {}
    return {
        "detected": True,
        "narrative": first(pairs, "事件經過", "經過"),
        "image_main": {
            "caption": first(pairs, "圖說"),
            "source": first(pairs, "圖片來源", "圖源"),
            "why": first(pairs, "為什麼選"),
        },
        "impact": first(pairs, "事件影響", "影響"),
        "reflections": {
            "q1": first(pairs, "省思一", "個人省思"),
            "q2": first(pairs, "省思二"),
            "q3": first(pairs, "省思三"),
        },
    }


def parse_one(path: Path) -> dict:
    doc = Document(str(path))
    table_text = []
    for table in doc.tables:
        for row in table.rows:
            table_text.extend(cell.text for cell in row.cells)
    full_text = "\n".join([p.text for p in doc.paragraphs] + table_text)
    pairs = iter_table_pairs(doc)

    class_name = first(pairs, "班級")
    seat = first(pairs, "座號")
    name = first(pairs, "姓名")
    event_text = first(pairs, "事件選擇")
    if not event_text:
        event_text = "\n".join([
            first(pairs, "事實句"),
            first(pairs, "差異句"),
            first(pairs, "省思句"),
            first(pairs, "地點"),
        ])
    event = detect_event(event_text)

    student = {
        "id": build_id(class_name, seat, name, path.stem),
        "class_name": class_name,
        "seat": seat,
        "name": name,
        "event": event,
        "facts": {
            "time": first(pairs, "時間"),
            "place": first(pairs, "地點"),
            "ethnic": first(pairs, "族群"),
            "people": first(pairs, "人物"),
        },
        "ai_process": {
            "question": first(pairs, "我問NotebookLM的問題", "我問 NotebookLM 的問題"),
            "ai_points": split_lines(first(pairs, "NotebookLM給我的3個重點", "NotebookLM 給我的 3 個重點")),
            "keywords": split_lines(first(pairs, "我抓到的3個關鍵詞", "我抓到的 3 個關鍵詞")),
            "original_ai_sentence": first(pairs, "AI原句", "AI 原句"),
            "revised_sentence": first(pairs, "我的改寫"),
            "revision_reason": first(pairs, "我修改的理由"),
            "process_log": first(pairs, "我怎麼判斷AI哪裡需要改", "我怎麼判斷 AI 哪裡需要改"),
            "verify_log": first(pairs, "我最後保留", "刪除", "查證"),
        },
        "perspectives": {
            "official": first(pairs, "官方說法"),
            "tribal": first(pairs, "部落／族人說法", "族人說法"),
            "difference": first(pairs, "最大的差異", "兩種說法最大的差異"),
        },
        "landscape": {
            "types": split_lines(first(pairs, "地景類型")),
            "looks_like_now": first(pairs, "現在看起來像什麼"),
            "after_knowing_history": first(pairs, "知道歷史後"),
        },
        "showcase_sentences": {
            "fact": first(pairs, "事實句"),
            "difference": first(pairs, "差異句"),
            "reflection": first(pairs, "省思句"),
        },
        "sources": parse_sources(first(pairs, "來源1", "來源 1"), first(pairs, "來源2", "來源 2"), first(pairs, "來源3", "來源 3")),
        "peer_review": {
            "checks": split_lines(first(pairs, "三個勾選", "三個勾")),
            "revised_sentence": first(pairs, "我建議同學改成這一句"),
        },
        "grading": {
            "basic": score_value(first(pairs, "基本事實")),
            "ai_use": score_value(first(pairs, "AI使用", "AI 使用")),
            "perspective": score_value(first(pairs, "立場比較")),
            "showcase": score_value(first(pairs, "展示句")),
        },
        "source_file": path.name,
    }

    legacy = legacy_supplement_if_present(full_text, pairs)
    if legacy:
        student["legacy_supplement"] = legacy
    return student


def main():
    sys.stdout.reconfigure(encoding="utf-8")
    SUBMIT_DIR.mkdir(parents=True, exist_ok=True)
    files = sorted(p for p in SUBMIT_DIR.glob("*.docx") if not p.name.startswith("~$"))
    students = []
    for path in files:
        try:
            st = parse_one(path)
            students.append(st)
            print(f"[OK] {path.name} -> {st['id']} ({st['event'] or '未判定'})")
        except Exception as exc:
            print(f"[ERR] {path.name} -> {exc}")

    payload = {
        "meta": {
            "school": "花蓮高商",
            "course": "縱谷無言．原住民族重大歷史事件",
            "unit": "原住民族重大歷史事件理解卡",
            "generated_at": datetime.now().isoformat(timespec="seconds"),
            "student_count": len(students),
            "schema": "student_card",
        },
        "students": students,
    }
    OUTPUT_JSON.write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")
    print(f"wrote {OUTPUT_JSON.relative_to(ROOT)} ({len(students)} students)")


if __name__ == "__main__":
    main()
